"""Study guide Word document generator."""
from __future__ import annotations

from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.shared import Pt

from exam_study_guide.config import SubjectConfig
from exam_study_guide.generators.word_utils import (
    add_centered_title,
    add_paragraph,
)
from exam_study_guide.matcher.review_workflow import MatchReview
from exam_study_guide.parsers.base import Question


def _make_explanation(q: Question, topic: str) -> str:
    """Generate a brief explanation for a question."""
    if q.qtype == "选择":
        if q.answer and q.options:
            correct_text = ""
            for opt in q.options:
                if opt.startswith(q.answer + "."):
                    correct_text = opt.split(".", 1)[1].strip()
                    break
            if correct_text:
                return f"【答案】{q.answer}\n【解析】本题考查“{topic}”。正确选项为 {q.answer}：{correct_text}。"
            return f"【答案】{q.answer}\n【解析】本题考查“{topic}”，请结合相关概念判断各选项。"
        return f"【答案】{q.answer or '（待补充）'}\n【解析】本题考查“{topic}”。"
    if q.qtype == "判断":
        base = f"【答案】{q.answer}\n【解析】本题考查“{topic}”。"
        return base + ("该命题表述存在错误。" if q.answer == "×" else "该命题表述正确。")
    if q.qtype == "简答":
        return f"【答题要点】本题属于“{topic}”相关简答题。回答时应条理清晰、简明扼要，突出关键词。"
    if q.qtype == "计算要点":
        return f"【复习要点】{q.stem}\n建议结合教材例题动手演算，计算过程要完整写出。"
    return ""


def generate_study_guide(
    config: SubjectConfig,
    questions: List[Question],
    review: MatchReview,
    output_path: Path,
) -> None:
    """Generate a Word study guide from a reviewed match."""
    doc = Document()
    add_centered_title(doc, f"《{config.subject}》考点复习指南", level=0, size=22)
    add_paragraph(doc, f"——{config.description}", style="Subtitle")
    add_paragraph(
        doc,
        "说明：本指南将考试大纲的每个考点与题库中相关题目进行匹配，并给出学习指导、"
        "参考答案与简要解析。计算题请结合教材例题动手演算。",
    )

    doc.add_page_break()

    # Table of contents
    add_paragraph(doc, "目录", style="Heading 1")
    for ch in config.chapters:
        add_paragraph(doc, ch.title, style="List Bullet")
        for point in ch.points:
            add_paragraph(doc, f"  {point.title}", style="List Bullet 2")
    doc.add_page_break()

    question_by_id: Dict[str, Question] = {q.qid: q for q in questions}

    for ch in config.chapters:
        add_paragraph(doc, ch.title, style="Heading 1")
        for idx, point in enumerate(ch.points, 1):
            add_paragraph(doc, f"{idx}. {point.title}", style="Heading 2")

            p = add_paragraph(doc, f"【学习指导】{point.guidance}")
            p.paragraph_format.space_after = Pt(8)

            matches = review.matches.get(point.title, [])
            if not matches:
                add_paragraph(
                    doc,
                    "（本题库暂未匹配到直接相关题目，请回归教材重点复习。）",
                    color=(128, 128, 128),
                )
                continue

            add_paragraph(
                doc, "【相关题目】", bold=True, size=11, color=(0, 112, 192)
            )
            for item in matches:
                q = question_by_id.get(item["qid"])
                if not q:
                    continue
                add_paragraph(
                    doc,
                    f"{q.qid}（{q.qtype}）",
                    bold=True,
                    size=10,
                    color=(31, 78, 121),
                )
                add_paragraph(doc, q.stem)
                for opt in q.options:
                    add_paragraph(doc, opt, style="List Bullet 2")
                exp = _make_explanation(q, point.title.split("：")[0])
                for line in exp.split("\n"):
                    if line.strip():
                        add_paragraph(doc, line, size=10)
                doc.add_paragraph()
        doc.add_page_break()

    doc.save(output_path)
