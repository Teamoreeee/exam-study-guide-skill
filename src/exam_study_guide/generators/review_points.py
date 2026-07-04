"""Review points + real-life examples Word document generator."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from exam_study_guide.config import SubjectConfig
from exam_study_guide.generators.word_utils import (
    add_centered_title,
    add_paragraph,
)


def generate_review_points(config: SubjectConfig, output_path: Path) -> None:
    """Generate a Word document with study points and real-life examples."""
    doc = Document()
    add_centered_title(doc, f"《{config.subject}》复习要点与生活实例", level=0, size=22)
    add_paragraph(
        doc,
        "本文档把每个考点的核心概念用生活化的例子再解释一遍，帮助你建立直觉。",
        style="Subtitle",
    )
    doc.add_paragraph()

    for ch in config.chapters:
        add_paragraph(doc, ch.title, style="Heading 1")
        for idx, point in enumerate(ch.points, 1):
            add_paragraph(doc, f"{idx}. {point.title}", style="Heading 2")

            p_guidance = add_paragraph(doc, point.guidance, size=11)
            p_guidance.paragraph_format.space_after = Pt(6)

            if point.examples:
                for example in point.examples:
                    p_example = add_paragraph(doc, f"生活例子：{example}", size=11)
                    p_example.paragraph_format.left_indent = Inches(0.2)
                    for run in p_example.runs:
                        run.font.color.rgb = RGBColor(0, 112, 192)
                    p_example.paragraph_format.space_after = Pt(8)
            else:
                add_paragraph(
                    doc,
                    "（暂无生活例子，建议结合教材例题理解。）",
                    size=10,
                    color=(128, 128, 128),
                )
        doc.add_page_break()

    doc.save(output_path)
