"""Shared helpers for Word document generation."""
from __future__ import annotations

from typing import Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DEFAULT_FONT = "Microsoft YaHei"


def set_run_font(
    run,
    name: str = DEFAULT_FONT,
    size: int | None = None,
    bold: bool = False,
    color: Tuple[int, int, int] | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(
    doc: Document,
    text: str,
    style: str | None = None,
    size: int | None = None,
    bold: bool = False,
    color: Tuple[int, int, int] | None = None,
):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_centered_title(doc: Document, text: str, level: int = 0, size: int = 22) -> None:
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        set_run_font(run, size=size, bold=True)
